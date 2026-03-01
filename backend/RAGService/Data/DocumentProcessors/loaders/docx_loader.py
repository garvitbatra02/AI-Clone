"""
DOCX Document Loader

Handles Microsoft Word (.docx) files.
Emits heading_hierarchy metadata for downstream semantic chunking.
"""

from __future__ import annotations

import asyncio
from pathlib import Path
from typing import Any, BinaryIO, Dict, List, Optional, Union

from RAGService.Data.DocumentProcessors.base import (
    BaseDocumentLoader,
    ProcessedDocument,
    SupportedFileType,
)


class DOCXLoader(BaseDocumentLoader):
    """
    Loader for Microsoft Word (.docx) files.
    
    Extracts text from paragraphs and optionally tables.
    Detects heading styles (Heading 1, Heading 2, etc.) and emits
    a heading_hierarchy in metadata for downstream semantic chunking.
    
    Emits structural metadata:
        - heading_hierarchy: List of {level, title, start_char} dicts
        - paragraph_count: Total paragraph count
        - table_count: Total table count
        - author, title, subject, created, modified (from core properties)
    
    Note: Requires python-docx package: pip install python-docx
    """
    
    def __init__(
        self,
        extract_tables: bool = True,
        preserve_formatting: bool = False
    ):
        """
        Initialize the DOCX loader.
        
        Args:
            extract_tables: Whether to extract text from tables
            preserve_formatting: Whether to preserve bold/italic markers
        """
        self.extract_tables = extract_tables
        self.preserve_formatting = preserve_formatting
    
    @property
    def supported_extensions(self) -> List[str]:
        return ["docx"]
    
    def _extract_text(self, source: Union[str, Path, BinaryIO]) -> tuple[str, Dict[str, Any]]:
        """
        Extract text from DOCX file with structural metadata.
        
        Detects heading styles and tracks their character offsets.
        Tables are interleaved at their actual document position
        (not appended at the end) and tagged with [TABLE] markers.
        """
        try:
            from docx import Document
            from docx.oxml.ns import qn
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX loading. "
                "Install with: pip install python-docx"
            )
        
        doc = Document(source)
        
        # Build an interleaved list of elements (paragraphs + tables)
        # by walking the document body XML in order
        content_parts = []
        heading_hierarchy = []  # {level, title, start_char}
        current_offset = 0
        
        body = doc.element.body
        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            
            if tag == "p":
                # It's a paragraph
                from docx.text.paragraph import Paragraph
                para = Paragraph(child, doc)
                text = para.text.strip()
                if not text:
                    continue
                
                # Check if it's a heading
                style_name = para.style.name if para.style else ""
                heading_level = self._get_heading_level(style_name)
                
                if heading_level is not None:
                    heading_hierarchy.append({
                        "level": heading_level,
                        "title": text,
                        "start_char": current_offset,
                    })
                
                # Apply formatting if requested
                if self.preserve_formatting:
                    formatted = []
                    for run in para.runs:
                        run_text = run.text
                        if run.bold:
                            run_text = f"**{run_text}**"
                        if run.italic:
                            run_text = f"*{run_text}*"
                        formatted.append(run_text)
                    text = "".join(formatted).strip()
                    if not text:
                        continue
                
                content_parts.append(text)
                current_offset += len(text) + 2  # +2 for "\n\n" join
            
            elif tag == "tbl" and self.extract_tables:
                # It's a table — interleave at actual position
                from docx.table import Table
                table = Table(child, doc)
                table_lines = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    table_lines.append(" | ".join(cells))
                if table_lines:
                    table_text = f"[TABLE]\n" + "\n".join(table_lines) + f"\n[/TABLE]"
                    content_parts.append(table_text)
                    current_offset += len(table_text) + 2
        
        content = "\n\n".join(content_parts)
        
        # Extract metadata
        metadata = {}
        if doc.core_properties:
            props = doc.core_properties
            if props.author:
                metadata["author"] = props.author
            if props.title:
                metadata["title"] = props.title
            if props.subject:
                metadata["subject"] = props.subject
            if props.created:
                metadata["created"] = str(props.created)
            if props.modified:
                metadata["modified"] = str(props.modified)
        
        metadata["paragraph_count"] = len(doc.paragraphs)
        metadata["table_count"] = len(doc.tables)
        metadata["heading_hierarchy"] = heading_hierarchy
        
        return content, metadata
    
    @staticmethod
    def _get_heading_level(style_name: str) -> Optional[int]:
        """
        Extract heading level from a paragraph style name.
        
        Returns:
            Heading level (1-9) or None if not a heading.
        """
        if not style_name:
            return None
        
        # python-docx style names: "Heading 1", "Heading 2", etc.
        lower = style_name.lower().strip()
        if lower.startswith("heading"):
            parts = lower.split()
            if len(parts) >= 2 and parts[-1].isdigit():
                return int(parts[-1])
            # "Heading" without a number — treat as level 1
            if lower == "heading":
                return 1
        
        # Also handle "Title" as a top-level heading
        if lower == "title":
            return 0
        
        return None
    
    def load(
        self,
        source: Union[str, Path, BinaryIO],
        metadata: Optional[Dict[str, Any]] = None
    ) -> ProcessedDocument:
        """Load a DOCX file."""
        meta = metadata or {}
        
        content, docx_meta = self._extract_text(source)
        meta.update(docx_meta)
        
        if isinstance(source, (str, Path)):
            source_str = str(Path(source).absolute())
        else:
            source_str = meta.get("source", "unknown")
        
        return ProcessedDocument(
            content=content,
            metadata=meta,
            source=source_str,
            file_type=SupportedFileType.DOCX
        )
    
    async def async_load(
        self,
        source: Union[str, Path, BinaryIO],
        metadata: Optional[Dict[str, Any]] = None
    ) -> ProcessedDocument:
        """
        Async version of load.
        
        Note: DOCX parsing is CPU-bound, so we run it in a thread pool.
        """
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None,
            lambda: self.load(source, metadata)
        )
